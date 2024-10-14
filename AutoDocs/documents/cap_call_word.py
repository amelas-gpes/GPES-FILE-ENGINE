from docx import Document
from docx.shared import Pt, Inches

from documents.parse_excel import *

import pandas as pd 
import os

from docx2pdf import convert
from docx.shared import RGBColor


def change_text(doc, position, message, font="Arial", size=10):
    # Access the specific paragraph you want to replace (e.g., the 3rd paragraph)
    paragraph = doc.paragraphs[position]

    # Clear any existing runs in the paragraph
    paragraph.clear()  # Clears the content and runs in the paragraph

    # Add a new run with the desired text
    new_run = paragraph.add_run(message)

    # Set font and size for the new run
    new_run.font.name = font  # Set font to Arial
    new_run.font.size = Pt(size)   # Set font size to 12 points

def show_doc_elements(doc):
    # Inspect paragraphs and runs
    for i, para in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: {para.text}")
        for j, run in enumerate(para.runs):
            print(f"  Run {j}: {run.text}")

    # Inspect tables
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            for cell in row.cells:
                print(f"  Cell: {cell.text}")

def update_field(doc, field, position, replacements):
    message = field.split()

    #Change field1
    modified_message = []
    i = 0
    endings = [".", ",", "!", "?", "..."]
    while i < len(message):
        ending = ""
        if message[i][-1] in endings:
            ending = message[i][-1]
            message[i] = message[i][:-1]

        if message[i] in replacements:
            if ending != "":
                modified_message.append(str(replacements[message[i]]) + ending)
            else:
                modified_message.append(str(replacements[message[i]]))
            i += 1
        else:
            if ending != "":
                modified_message.append(message[i] + ending)
            else:
                modified_message.append(message[i])
            i += 1
    
    modified_message = " ".join(modified_message)

    change_text(doc, position, modified_message)

def create_cap_call_pdf(doc, excel, fund_info, inv_info, total_fund_info, output_path, logo_path, text_fields, start_delim, end_delim, cap_call_table_data):
    wire_instructions = {
        "Bank Name:" : "Chase",
        "Bank Address:" : "277 Park Ave New York, NY 10172",
        "ABA Number:" : "021000021 (Domestic Wires)",
        "Account Name:" : fund_info["Partner Name"],
        "Account Number:" : "932859662",
        "SWIFT Code:" : "CHASUS33 (International Wires)"
    }

    #add logos to pages
    for section in doc.sections:
        header = section.header
        
        header_paragraph = header.paragraphs[0]
        header_paragraph.clear()

        header_paragraph.add_run().add_picture(logo_path, width=Inches(1.5))  # Adjust the size as needed

    #change first page header
    header = doc.tables[0]

    header.rows[2].cells[1].text = fund_info["Partner Name"]
    header.rows[4].cells[1].text = fund_info["Re"]
    header.rows[6].cells[1].text = fund_info["Notice Date"]

    header.rows[8].cells[1].text = ""
    header_fund_info = header.rows[8].cells[1].paragraphs[0].add_run(fund_info["Due Date"])
    header_fund_info.bold = True
    header_fund_info.font.name = "Arial"
    header_fund_info.font.size = Pt(10)

    #Change fields
    replacements = {
        "<Section#>" : "3.3",
        "<Notice_Date>" : fund_info["Notice Date"],
        "<Total_Amount_Due>" : f"${int(fund_info['Net Amount Due / (Payable)']):,}",
        "<Due_Date>" : fund_info["Due Date"],
        "<Contact_Name>" : fund_info["Contact Name"],
        "<Contact_Email>" : fund_info["Contact Email"]
    }

    update_field(doc, text_fields[0], 3, replacements)
    update_field(doc, text_fields[1], 5, replacements)
    update_field(doc, text_fields[2], 7, replacements)
    update_field(doc, text_fields[3], 10, replacements)

    
    instructions = doc.tables[1]

    #Fill in wire instructions table
    for row in instructions.rows:
        if row.cells[0].text.strip():
            if str(row.cells[0].text.strip()) in wire_instructions:
                row.cells[1].text = wire_instructions[str(row.cells[0].text.strip())]


    change_text(doc, 15, fund_info["Partner Name"])


    re = fund_info["Re"].split()
    re = ' '.join(re[:3])
    change_text(doc, 21, "Re: " + re)


    table = doc.tables[2]

    #Empty table
    for row in table.rows:
        row._element.getparent().remove(row._element)
    
    capital_call_data = [
        ["", "", "Total Fund", "", "", "Your Share"],
        ["Capital Call", "", "", "", "", ""]
    ]

    for i in cap_call_table_data:
        if i[1][1:-1] not in total_fund_info: continue
        capital_call_data.append([i[0], "$", total_fund_info[i[1][1:-1]], "", "$", inv_info[i[1][1:-1]]])

    """
    #add investments to capital call table
    for key in inv_info:
        if "Investment" == key.split()[0]:
            if total_fund_info[key] == "nan": continue
            capital_call_data.append([key, "$", total_fund_info[key], "", "$", inv_info[key]])
    

    #add management fees to capital call table
    capital_call_data.append(["Management Fees", "", total_fund_info["Gross Mgmt Fee"], "", "", inv_info["Gross Mgmt Fee"]])

    #add partnership expenses to capital call table
    capital_call_data.append(["Partnership Expenses", "", total_fund_info["Pshp Exp"], "", "", inv_info["Pshp Exp"]])
    """

    #add ending space
    capital_call_data.append(["", "", "", "", "", ""])

    # Insert rows to capital call table
    for data_row in capital_call_data:
        row_cells = table.add_row().cells
        for i, data in enumerate(data_row):
            row_cells[i].text = data


    additional_information_data = [
        ["Additional Information", "", "", "", "", ""]
    ]
    

    # Insert rows to additional information table
    for data_row in additional_information_data:
        row_cells = table.add_row().cells
        for i, data in enumerate(data_row):
            row_cells[i].text = data
    
    
    header.rows[0].cells[1].text = ""
    header_inv_name = header.rows[0].cells[1].paragraphs[0].add_run(inv_info["Partner Name"])
    header_inv_name.bold = True
    change_text(doc, 20, "Investor: " + inv_info["Partner Name"])
    #Fill in the info
    instructions.rows[7].cells[1].text = inv_info["Partner Name"]

    """
    #Fill in investor info table
    for row in table.rows:
        if row.cells[0].text.strip():
            if str(row.cells[0].text.strip()) in inv_info:
                row.cells[5].text = "{:,}".format(inv_info[str(row.cells[0].text.strip())])
    """

    # add delimiters
    # Create the new paragraph
    new_paragraph = doc.add_paragraph()

    # Add the text to the paragraph
    run = new_paragraph.add_run(f"{start_delim} {end_delim}")

    # Set the text color to white (RGB for white is 255, 255, 255)
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Move the newly added paragraph to the desired index
    doc._body._element.insert(7, new_paragraph._element)

    print(output_path)
    investor_docx = output_path + ".docx"
    investor_pdf = output_path + ".pdf"
    doc.save(investor_docx)
    # Convert the Word document to PDF
    convert(investor_docx, investor_pdf)

    os.remove(investor_docx)

